# -*- coding: utf-8 -*-

from selenium.webdriver.common.by import By
from selenium import webdriver
import time
import random
import os
from urllib import request
import docx
from docx.shared import Pt
from docx.shared import Inches
from selenium.webdriver import ChromeOptions
from docx.oxml.ns import qn
from PIL import Image
import pytesseract

option = ChromeOptions()
# 隐藏被浏览器被自动化工具控制的提示
option.add_experimental_option('excludeSwitches', ['enable-automation'])
driver = webdriver.Chrome(options=option)
driver.maximize_window()  # 浏览器窗口最大化


class JianBiaoSpider(object):
    # 将Tesseract软件的安装及存储路径添加进来
    pytesseract.pytesseract.tesseract_cmd = r'/usr/local/bin/tesseract'
    # tessdata_dir_config = r'--tessdata-dir "F:\Program Files\Tesseract-OCR\tessdata"'

    def __init__(self, title):
        
        
        
        # 待爬取规范的网址首页，这里爬的是《混凝土结构设计规范》
        self.ini_url = 'http://www.jianbiaoku.com/webarbs/book/209/2438396.shtml'
        self.title = title  # 规范名称
        self.file = docx.Document()  # 创建内存中的word文档对象
        self.file.add_heading(self.title, 2)  # 为文档添加标题
        # 设置规范中图片的本地存储路径，如没有该文件夹则创建
        self.image_path = os.path.join(os.path.dirname(__file__), 'image')
        is_exists = os.path.exists(self.image_path)
        if not is_exists:
            os.makedirs(self.image_path)
        # 设置爬取结果的本地存储路径，如没有该文件夹则创建
        self.file_path = os.path.join(os.path.dirname(__file__), '规范文件')
        is_exists = os.path.exists(self.file_path)
        if not is_exists:
            os.makedirs(self.file_path)

    def get_text(self, url):
        # 初始化刷新重试次数
        count1 = count2 = 0
        while True:
            try:
                driver.get(url)  # 加载网址内容
                time.sleep(random.choice([3, 1.3, 1, 1.5, 0.8, 0.6]))  # 设置随机休眠时间
                # 死循环用于判断是否已摆脱安全验证页面，如已加载到正文页面则跳出循环
                while True:
                    # 判断是否跳出安全验证页面，如跳出则执行验证码OCR自动识别，如已加载到正文页面则跳出循环
                    if driver.find_elements(By.XPATH, '//div[@class="access-validate-wrap"]/div[@class="tips-line"]'):
                        print('进入验证')
                        driver.get_screenshot_as_file('window1.png')
                        image = Image.open('window1.png')
                        # 采用截屏+抠图方式获取验证码区域，该处抠图区域坐标（左上,右下）因电脑分辨率和缩放比例不同会有差异，需自行调试
                        captcha = image.crop((1060, 710, 1200, 780))
                        captcha.save('captcha1.png')  # 保存目标验证码抠图图片
                        image = Image.open('captcha1.png')  # 加载目标验证码图片
                        text = pytesseract.image_to_string(image, lang='eng')  # OCR识别
                        yzm = ''
                        # 保留非空字符并添加到yzm中保存
                        for num in text:
                            if num != ' ' and num != '\n':
                                yzm = yzm + num
                        print(yzm)
                        # 定位到验证码输入框并输入验证码识别结果
                        driver.find_element(By.XPATH, 'ipt_captcha').send_keys(yzm)
                        time.sleep(1)
                        # 定位到验证码提交按钮并点击
                        driver.find_element(By.XPATH, 'btn_captcha_confirm').click()
                        count1 += 1
                        if count1 > 4:
                            driver.refresh()
                        time.sleep(1)
                    else:
                        # 已加载到正文页面则跳出循环
                        print('跳出验证')
                        break

                while True:
                    print('等待内容加载')
                    time.sleep(1)
                    # 定位所有的p标签，正文文本和图片全部在p标签下
                    li_list = driver.find_elements(By.XPATH,
                                                   '//div[@class="book_right"]/div[@class="book_content"]//p')
                    if len(li_list) == 0:
                        continue
                    elif len(li_list) == 1:
                        cont = li_list[0].text.strip()
                        if cont.startswith('页面正在加载中'):
                            print('等待数据加载...')
                            continue
                        break
                    else:
                        break
                print('标签: ', li_list)
                for li in li_list:
                    # 判断p标签下是否存在图片
                    if li.find_elements(By.XPATH, './/img'):
                        # 找到p标签的所有子标签
                        son_tags = li.find_elements(By.XPATH, './*')
                        # 若p标签下存在图片，且p标签的子标签个数大于1，则说明图片嵌入文字中（即用小图片表示特殊字符，如带有上下角标的希腊字母，与文字在同一行中）
                        if len(son_tags) > 1:
                            # 将文字和图片依次写入word文档中，并设置字体，图片尺寸等
                            run_cont = self.file.add_paragraph().add_run()
                            self.file.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
                            self.file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),
                                                                               u'宋体')  # 设置中文字体使用字体2->宋体
                            run_cont.font.size = Pt(10)  # 设置字体大小
                            for son_tag in son_tags:
                                cont = son_tag.text.strip()  # 文本去除前后空格
                                if cont:
                                    run_cont.add_text(cont)  # 文本写入word
                                    print(cont)
                                if son_tag.get_attribute('src'):
                                    img_url = son_tag.get_attribute('src')  # 定位至图片的url链接
                                    image_name = img_url.split('/')[-1]  # 提取图片标识名称
                                    print(image_name, img_url)
                                    img_path = os.path.join(self.image_path, image_name)  # 保存图片至先前建好的存放目录
                                    request.urlretrieve(img_url, img_path)  # 下载图片
                                    run_cont.add_picture(img_path, width=Inches(0.1))  # 写入图片至word，并设置尺寸
                        # 若p标签下存在图片，且p标签的子标签个数等于1，则说明该p标签是一张独立的大图，单独一行放置
                        else:
                            cont = son_tags[0].text.strip()
                            run_cont = self.file.add_paragraph().add_run()
                            self.file.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
                            self.file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),
                                                                               u'宋体')  # 设置中文字体使用字体2->宋体
                            run_cont.font.size = Pt(10)
                            run_cont.add_text(cont)
                            print(cont)
                            if son_tags[0].get_attribute('src'):
                                img_url = son_tags[0].get_attribute('src')
                                image_name = img_url.split('/')[-1]
                                print(image_name, img_url)
                                img_path = os.path.join(self.image_path, image_name)
                                request.urlretrieve(img_url, img_path)
                                self.file.add_picture(img_path, width=Pt(400))  # 写入图片至word，并设置尺寸
                    else:
                        # 如果p标签下不存在图片，则提取所有的文字，保存至word中
                        cont = li.text.strip()
                        self.file.add_paragraph(cont)
                        print(cont, 'aaaaa')
                break

            # 若页面超时无响应，则刷新页面，并限制次数
            except Exception as e:
                print(e)
                count2 += 1
                if count2 > 4:
                    break
                driver.refresh()

        # 每加载一页保存一次
        print('保存一页', self.title)
        self.file.save("{}/{}.docx".format(self.file_path, self.title))
        # 定位至下一页按钮的title
        next_cate = driver.find_element(By.XPATH, '//div[@class="next_catalog"]/a').get_attribute('title')
        self.file.add_page_break()  # 写入完成一页内容后，插入分页
        # 若下一页没有了，则结束，若下一页还存在，则继续加载下一页内容
        if next_cate != '下一章：没有了':
            # 定位至下一页的url链接
            next_url = driver.find_element(By.XPATH, '//div[@class="next_catalog"]/a').get_attribute('href')
            # print(next_url)
            time.sleep(random.choice([3, 1.3, 1, 1.5, 0.8, 0.6]))  # 设置随机休眠时间
            self.get_text(next_url)  # 加载下一页

    def main(self):
        # 主函数
        self.get_text(self.ini_url)


if __name__ == '__main__':
    
    # 修改爬取的规范位置
    bz1 = JianBiaoSpider('混凝土结构设计规范')
    bz1.main()
